"""
تبدیل فایل Word (.docx) به HTML با حفظ فرمت (بولد، ایتالیک، آندرلاین، ترازبندی و ...).

کاربرد: کاربر می‌تواند به‌جای تایپ متن، فایل ورد را ارسال کند.
متن استخراج‌شده عیناً با فرمت اصلی در ادیتور سامانه قضایی وارد می‌شود.

⚠️ اصلاحیهٔ مهم:
قبلاً `from docx import Document` در سطح ماژول بود؛ اگر بستهٔ python-docx
نصب نبود، کل ماژول (و همهٔ هندلرهایی که آن را import می‌کردند) با خطای:
    ModuleNotFoundError: No module named 'docx'
کرش می‌کرد — نتیجه: ربات به فایل ورد «هیچ واکنشی نشان نمی‌داد» و وارد
مرحلهٔ بعد نمی‌شد (گزارش کارفرما).

حالا import داخل try/except است و در صورت نبودِ python-docx، یک پارسر
جایگزینِ کاملاً stdlib (zipfile + xml.etree) با همان خروجی استفاده
می‌شود؛ یعنی docx_to_html / docx_to_plain_text همیشه کار می‌کنند.
"""

import os
import logging

logger = logging.getLogger(__name__)

# حداکثر حجم مجاز فایل ورد (بایت)
MAX_DOCX_SIZE = 10 * 1024 * 1024  # 10 MB

# ── ایمپورت امن python-docx ────────────────────────────────────────────────
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _HAS_PYTHON_DOCX = True
except ImportError:
    _HAS_PYTHON_DOCX = False
    import zipfile
    import xml.etree.ElementTree as ET
    logger.warning(
        "بستهٔ python-docx نصب نیست — از پارسر جایگزین (zipfile+XML) استفاده "
        "می‌شود. برای کیفیت بهتر نصب کنید: pip install python-docx")


# ═════════════════════════════════════════════════════════════════════════════
# پارسر جایگزین (stdlib) — خروجی مثل python-docx؛ همیشه تعریف می‌شود تا
# در صورت خطای زمان اجرا هم قابل استفاده باشد
# ═════════════════════════════════════════════════════════════════════════════
import zipfile
import xml.etree.ElementTree as ET

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = "{%s}" % _W_NS

class _FBLength:
    """شبیه docx.shared.Length — فقط پراپرتی pt."""
    def __init__(self, half_points: int):
        self._half_points = half_points

    @property
    def pt(self) -> float:
        return self._half_points / 2.0

class _FBFont:
    """شبیه docx.text.font.Font — فقط strike و size (استفاده‌شده در _run_to_html)."""
    def __init__(self, strike: bool, size_half_pt):
        self.strike = strike
        self.size = (_FBLength(size_half_pt) if size_half_pt is not None else None)

class _FBRun:
    """شبیه docx.text.run.Run."""
    def __init__(self, text: str, bold: bool, italic: bool,
                 underline: bool, strike: bool, size_half_pt):
        self.text = text
        self.bold = bold
        self.italic = italic
        self.underline = underline
        self.font = _FBFont(strike, size_half_pt)

class _FBParagraph:
    """شبیه docx.text.paragraph.Paragraph — alignment مقدار متنی است
    (right/center/left/both) که _align_to_css پشتیبانی می‌کند."""
    def __init__(self, runs, alignment):
        self.runs = runs
        self.alignment = alignment

    @property
    def text(self) -> str:
        return "".join(r.text for r in self.runs)

class _FBDoc:
    """شبیه docx.Document — فقط paragraphs."""
    def __init__(self, paragraphs):
        self.paragraphs = paragraphs

def _fb_parse_align(pPr) -> str:
    """مقدار w:jc را برمی‌گرداند (رشتهٔ خام مثل right/center/left/both)."""
    if pPr is None:
        return None
    jc = pPr.find(_W + "jc")
    if jc is None:
        return None
    return jc.get(_W + "val")


def _fb_parse_run(r):
    rPr = r.find(_W + "rPr")
    bold = italic = underline = strike = False
    size_half_pt = None
    if rPr is not None:
        bold = rPr.find(_W + "b") is not None
        italic = rPr.find(_W + "i") is not None
        underline = rPr.find(_W + "u") is not None
        strike = rPr.find(_W + "strike") is not None
        sz = rPr.find(_W + "sz")
        if sz is not None:
            v = sz.get(_W + "val")
            if v and v.isdigit():
                size_half_pt = int(v)

    # متن run — w:t (و w:tab/w:br برای حفظ ساختار خط)
    parts = []
    for child in r:
        if child.tag == _W + "t":
            parts.append(child.text or "")
        elif child.tag == _W + "tab":
            parts.append("\t")
        elif child.tag == _W + "br":
            parts.append("\n")
    text = "".join(parts)
    if not text:
        return None
    return _FBRun(text, bold, italic, underline, strike, size_half_pt)


def _fb_open_document(filepath: str) -> _FBDoc:
    """باز کردن .docx با zipfile و پارس word/document.xml.

    فقط پاراگراف‌های مستقیمِ بدنه (w:body → w:p) برگردانده می‌شوند —
    دقیقاً مثل doc.paragraphs در python-docx."""
    with zipfile.ZipFile(filepath) as z:
        xml_bytes = z.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    body = root.find(_W + "body")
    if body is None:
        p_elements = list(root.iter(_W + "p"))
    else:
        p_elements = [el for el in body if el.tag == _W + "p"]

    paragraphs = []
    for p in p_elements:
        pPr = p.find(_W + "pPr")
        alignment = _fb_parse_align(pPr)
        runs = []
        for r in p.findall(_W + "r"):
            fb_run = _fb_parse_run(r)
            if fb_run is not None:
                runs.append(fb_run)
        paragraphs.append(_FBParagraph(runs, alignment))
    return _FBDoc(paragraphs)


def _open_docx_document(filepath: str):
    """باز کردن سند ورد — python-docx در صورت وجود، وگرنه پارسر جایگزین.

    اگر python-docx هنگام باز کردن فایل خاصی خطا داد، به‌طور خودکار به
    پارسر جایگزین سوئیچ می‌شود تا ربات هرگز به‌خاطر یک فایل ورد متوقف
    نشود."""
    if _HAS_PYTHON_DOCX:
        try:
            return Document(filepath)
        except Exception as pydocx_err:
            logger.warning(
                f"python-docx نتوانست فایل را باز کند ({pydocx_err!r}) — "
                "ادامه با پارسر جایگزین")
    return _fb_open_document(filepath)


# ═════════════════════════════════════════════════════════════════════════════
# تبدیل به HTML
# ═════════════════════════════════════════════════════════════════════════════
def _align_to_css(alignment) -> str:
    """تبدیل ترازبندی ورد به استایل CSS.

    هم enum پایتون-داکس (name مثل RIGHT/CENTER/LEFT/JUSTIFY) و هم
    مقدار متنی پارسر جایگزین (right/center/left/both/start/end) پشتیبانی
    می‌شود."""
    if alignment is None:
        return "text-align: right;"  # پیش‌فرض: راست‌چین
    name = getattr(alignment, "name", None)
    val = (name if name else str(alignment)).strip().upper()
    if "RIGHT" in val or val == "END":
        return "text-align: right;"
    if "CENTER" in val:
        return "text-align: center;"
    if "LEFT" in val or val == "START":
        return "text-align: left;"
    if "JUSTIFY" in val or val == "BOTH":
        return "text-align: justify;"
    return "text-align: right;"  # پیش‌فرض: راست‌چین


def _run_to_html(run) -> str:
    """
    یک Run ورد را به HTML تبدیل می‌کند.
    بولد، ایتالیک، آندرلاین، خط‌خورده و اندازه فونت حفظ می‌شوند.
    """
    text = run.text or ""
    if not text:
        return ""

    # escape کاراکترهای HTML
    import html as html_lib
    text = html_lib.escape(text, quote=False)

    # حفظ فاصله‌های ابتدای خط
    if text.startswith(" "):
        leading = len(text) - len(text.lstrip(" "))
        text = ("\u00a0" * leading) + text[leading:]
    text = text.replace("  ", "\u00a0 ")

    # بررسی فرمت‌ها
    tags = []
    if run.bold:
        tags.append("b")
    if run.italic:
        tags.append("i")
    if run.underline:
        tags.append("u")
    try:
        if run.font and run.font.strike:
            tags.append("s")
    except Exception:
        pass

    #包裹 متن با تگ‌ها
    for tag in tags:
        text = f"<{tag}>{text}</{tag}>"

    # اندازه فونت
    try:
        font_size = run.font.size if run.font else None
    except Exception:
        font_size = None
    if font_size:
        # ورد اندازه را در half-points ذخیره می‌کند (1 pt = 2 half-pts)
        size_pt = font_size.pt
        if size_pt and size_pt != 12:  # 12pt پیش‌فرض است
            text = f'<span style="font-size: {size_pt}pt;">{text}</span>'

    return text


def _paragraph_to_html(paragraph) -> str:
    """
    یک پاراگراف ورد را به تگ <p> HTML تبدیل می‌کند.
    """
    # جمع‌آوری HTML تمام run‌های پاراگراف
    inner_html = ""
    for run in paragraph.runs:
        inner_html += _run_to_html(run)

    # اگر پاراگراف خالی است
    if not inner_html.strip():
        return "<p><br></p>"

    # ترازبندی
    align_style = _align_to_css(paragraph.alignment)

    return f'<p style="{align_style}">{inner_html}</p>'


def docx_to_html(filepath: str) -> str:
    """
    فایل .docx را می‌خواند و HTML مناسب برای ادیتور سامانه برمی‌گرداند.

    پارامترها:
        filepath: مسیر فایل .docx

    بازگشت:
        رشته HTML (مثلاً <p>متن <b>بولد</b></p><p>خط دوم</p>)
    """
    try:
        doc = _open_docx_document(filepath)
    except Exception as e:
        logger.error(f"خطا در باز کردن فایل ورد: {e}")
        raise ValueError(f"خطا در باز کردن فایل ورد: {e}")

    html_parts = []
    for para in doc.paragraphs:
        html_parts.append(_paragraph_to_html(para))

    result = "".join(html_parts)

    if not result.strip() or result == "<p><br></p>":
        return "<p><br></p>"

    return result


def docx_to_plain_text(filepath: str) -> str:
    """
    فایل .docx را می‌خواند و متن خام (بدون HTML) برمی‌گرداند.
    برای نمایش در پیش‌نمایش پیام‌رسان استفاده می‌شود.
    """
    try:
        doc = _open_docx_document(filepath)
    except Exception as e:
        logger.error(f"خطا در باز کردن فایل ورد: {e}")
        return ""

    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)

    return "\n".join(lines)


def validate_docx(filepath: str) -> tuple:
    """
    اعتبارسنجی فایل ورد.
    بازگشت: (is_valid: bool, error_message: str)
    """
    if not os.path.exists(filepath):
        return False, "فایل یافت نشد."

    if not filepath.lower().endswith(".docx"):
        return False, "فقط فایل‌های .docx پشتیبانی می‌شوند."

    size = os.path.getsize(filepath)
    if size > MAX_DOCX_SIZE:
        mb = size / (1024 * 1024)
        return False, f"حجم فایل ({mb:.1f} MB) بیش از حد مجاز (10 MB) است."

    if size < 100:
        return False, "فایل خالی یا ناقص به نظر می‌رسد."

    return True, ""


async def download_docx_from_bale(bot, file_id: str, user_id: int) -> str:
    """
    دانلود فایل ورد از بله و ذخیره محلی.
    بازگشت: مسیر فایل دانلودشده یا None در صورت خطا.
    """
    import time
    try:
        file_info = await bot.get_file(file_id)
        if not file_info.file_path:
            logger.error(f"[DOCX] مسیر فایل تلگرام خالی است برای {file_id}")
            return None

        filename = f"docx_{user_id}_{int(time.time()*1000)}.docx"
        await bot.download_file(file_info.file_path, filename)
        return filename
    except Exception as e:
        logger.error(f"[DOCX] خطا در دانلود فایل ورد: {e}")
        return None
